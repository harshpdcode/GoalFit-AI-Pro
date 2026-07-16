# -*- coding: utf-8 -*-
r"""
GoalFit AI Pro - Project Documentation Generator
Generates a fully formatted .docx file matching the reference document style.
Run:  venv\Scripts\python.exe generate_docs.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

DOC_PATH = "documentation/GoalFit_AI_Pro_Documentation.docx"

# ─── Colour palette (matching reference: dark blue headings, white body) ───
C_HEADING   = RGBColor(0x00, 0x33, 0x66)   # deep navy
C_SUBHEAD   = RGBColor(0x00, 0x55, 0x99)   # medium blue
C_ACCENT    = RGBColor(0x00, 0x96, 0x88)   # teal  (used for table headers)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_BG  = RGBColor(0xE8, 0xF4, 0xFF)   # light-blue row fill
C_BODY      = RGBColor(0x1A, 0x1A, 0x2E)   # near-black body text

def set_cell_bg(cell, hex_color):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = C_HEADING
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.bold = True
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = C_SUBHEAD
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = C_SUBHEAD
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    return p

def add_body(doc, text, italic=False, bold=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = C_BODY
        run.font.italic = italic
        run.font.bold   = bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = C_BODY
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = C_BODY
    return p

def header_table(doc, cols, rows_data, col_widths=None):
    """Creates a styled table with a teal header row."""
    table = doc.add_table(rows=1 + len(rows_data), cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
        set_cell_bg(hdr[i], '009688')
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold  = True
            run.font.color.rgb = C_WHITE
            run.font.size  = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri + 1].cells
        bg  = 'E8F4FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            row[ci].text = str(val)
            set_cell_bg(row[ci], bg)
            for run in row[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.color.rgb = C_BODY

    # Column widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return table


# ═══════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# Default paragraph font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


# ─────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJECT DOCUMENTATION")
run.font.size  = Pt(13)
run.font.bold  = True
run.font.color.rgb = C_BODY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GoalFit AI Pro")
run.font.size  = Pt(32)
run.font.bold  = True
run.font.color.rgb = C_HEADING

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Personalised Fitness & Nutrition Platform\nwith Professional Trainer / Dietician Marketplace")
run.font.size  = Pt(14)
run.font.color.rgb = C_SUBHEAD
run.font.italic = True

doc.add_paragraph("\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted by: Harsh Pandya\nDepartment of Computer Applications\nAcademic Year: 2025–2026")
run.font.size  = Pt(12)
run.font.color.rgb = C_BODY

doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# TABLE OF CONTENTS  (manual)
# ─────────────────────────────────────────────────────────────
add_h1(doc, "TABLE OF CONTENTS")
toc = [
    ("1",      "PROJECT TITLE",                          "1"),
    ("2",      "PROBLEM DEFINITION",                     "1"),
    ("3",      "NEED OF PROJECT",                        "2"),
    ("3.1",    "Current System and Its Drawbacks",       "3"),
    ("3.2",    "Proposed System and Its Features",       "4"),
    ("4",      "REQUIREMENTS",                           "5"),
    ("4.1",    "Software Requirements",                  "6"),
    ("4.2",    "Hardware Requirements",                  "7"),
    ("5",      "TIME DURATION",                          "8"),
    ("6",      "TECHNOLOGY USED",                        "9"),
    ("7",      "FUNCTIONAL AND NON-FUNCTIONAL REQUIREMENTS", "10"),
    ("8",      "DESIGN",                                 "11"),
    ("8.1",    "System Flow Diagram",                    "11"),
    ("8.2",    "UML / Data Flow Diagrams",               "12"),
    ("8.2.1",  "Use Case Diagram (Context / 0 Level)",   "13"),
    ("8.2.2",  "Activity Diagram (First Level)",         "14"),
    ("8.2.3",  "Sequence Diagram (Second Level)",        "15"),
    ("8.2.4",  "Class Diagram",                         "16"),
    ("8.3",    "ER Diagram",                             "17"),
    ("8.4",    "Data Dictionary",                        "18"),
    ("8.5",    "Module Screenshots",                     "20"),
    ("9",      "APPLICATION",                            "46"),
    ("10",     "EXPECTED OUTCOMES",                      "47"),
    ("11",     "FUTURE SCOPE",                           "48"),
    ("12",     "REFERENCES",                             "49"),
]
for sr, title, pg in toc:
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.5))
    run = p.add_run(f"{sr}   {title}")
    run.font.size = Pt(11)
    run.font.color.rgb = C_BODY
    run2 = p.add_run(f"\t{pg}")
    run2.font.size = Pt(11)
    run2.font.color.rgb = C_BODY

doc.add_page_break()


# ═════════════════════════════════════════════════════════════
# 1. PROJECT TITLE
# ═════════════════════════════════════════════════════════════
add_h1(doc, "1. PROJECT TITLE")
add_body(doc,
    "GoalFit AI Pro – An AI-Powered Personalised Fitness and Nutrition Platform "
    "with an Integrated Professional Trainer / Dietician Marketplace"
)
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 2. PROBLEM DEFINITION
# ═════════════════════════════════════════════════════════════
add_h1(doc, "2. PROBLEM DEFINITION")
add_body(doc,
    "In today's fast-paced lifestyle, maintaining a healthy body and mind has become "
    "increasingly difficult. A large portion of the population struggles to:"
)
bullets = [
    "Design a diet plan that is personalised to their age, weight, goal, and food preferences.",
    "Follow a structured workout regimen without professional guidance.",
    "Track daily calorie intake, water consumption, and body-weight progression.",
    "Find affordable, verified, and trustworthy fitness professionals (trainers/dieticians) "
    "in one place.",
    "Receive AI-driven insights and goal predictions without expensive consultations.",
]
for b in bullets:
    add_bullet(doc, b)

add_body(doc,
    "\nExisting applications are either too generic, too expensive, or lack integration between "
    "fitness tracking and professional coaching. There is no unified platform that combines "
    "AI-personalised planning, real-time tracking, and a vetted professional marketplace."
)
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 3. NEED OF PROJECT
# ═════════════════════════════════════════════════════════════
add_h1(doc, "3. NEED OF PROJECT")
add_body(doc,
    "GoalFit AI Pro addresses a critical gap in the fitness-technology ecosystem. The following "
    "factors highlight the necessity of this project:"
)
for b in [
    "Rapid growth in lifestyle diseases (obesity, diabetes, hypertension) driven by sedentary habits.",
    "Rising demand for personalised health guidance that is affordable and accessible.",
    "Lack of a single platform that combines AI recommendations, daily tracking, and live professional coaching.",
    "Need for trainers and dieticians to reach clients digitally and manage them through a structured SaaS dashboard.",
    "Requirement for downloadable meal/workout PDF plans for offline use.",
    "Growing interest in body-composition tracking with progress photos and weight logs.",
]:
    add_bullet(doc, b)
doc.add_paragraph()


# 3.1
add_h2(doc, "3.1 Current System and Its Drawbacks")
add_body(doc,
    "Most existing health & fitness apps operate in silos and suffer from several limitations:"
)
rows = [
    ("Generic Diet Plans",   "Apps provide one-size-fits-all meal plans with no personalisation by goal type, diet preference, or food culture."),
    ("No Professional Link", "Users cannot directly hire and interact with verified trainers or dieticians inside the same app."),
    ("Fragmented Tools",     "Calorie trackers, workout planners, progress loggers, and coaching platforms exist as separate apps."),
    ("No AI Predictions",    "Traditional apps lack AI-driven goal completion date predictions or step recommendations."),
    ("No PDF Export",        "Meal and workout plans cannot be exported as formatted PDFs for offline or printed use."),
    ("Limited Admin Control","No centralised admin panel to manage users, professionals, verify accounts, or view platform revenue."),
    ("No Marketplace",       "Users have no transparent way to browse, compare, and purchase plans from certified professionals."),
]
header_table(doc,
    ["Drawback", "Description"],
    rows,
    col_widths=[2.0, 4.5]
)
doc.add_paragraph()


# 3.2
add_h2(doc, "3.2 Proposed System and Its Features")
add_body(doc,
    "GoalFit AI Pro is a full-stack, AI-powered fitness platform built with Flask (Python), "
    "MySQL, and modern JavaScript. Key features of the proposed system:"
)
features = [
    ("AI-Personalised Diet & Workout Plans",
     "Plans are generated based on user health profile (age, gender, weight, goal, diet type). "
     "Both vegetarian, non-vegetarian, and vegan diets are supported."),
    ("BMI Calculator & Goal Predictor",
     "Computes BMI, categorises it, recommends daily steps, and predicts goal completion date "
     "using a weekly-change-rate algorithm."),
    ("Professional Marketplace",
     "Users can browse trainer/dietician cards with ratings and pricing, view detailed profiles "
     "with transformation stories, and hire professionals via Razorpay payment gateway."),
    ("Professional SaaS Dashboard (Pro Portal)",
     "Professionals get a dedicated dark-themed portal to manage clients, create custom meal/workout "
     "plans, track client progress, upload transformation stories, and view earnings."),
    ("Real-Time Progress Tracking",
     "Weight logs, progress photos (gallery with sharing controls), BMI history, and daily calorie "
     "tracking are all available."),
    ("Water Intake Tracker",
     "Tracks daily water glasses consumed with visual progress bars."),
    ("Chat System",
     "Built-in overlay chat between professional and client available directly from the client "
     "detail page."),
    ("PDF Report Generator",
     "Users can download a fully formatted PDF containing their personalised diet and workout plan."),
    ("Admin Panel",
     "Centralized admin dashboard to manage users, professionals (verify/suspend), view payments, "
     "revenue analytics, and reply to feedback."),
    ("Notifications System",
     "Real-time notification badges for hire requests, payments, and reviews on the professional portal."),
    ("Rate Limiting & Security",
     "Flask-Limiter applied globally; passwords hashed with Werkzeug; environment credentials via python-dotenv."),
]
for title, desc in features:
    p = doc.add_paragraph()
    run1 = p.add_run(f"• {title}: ")
    run1.font.bold  = True
    run1.font.size  = Pt(11)
    run1.font.color.rgb = C_SUBHEAD
    run2 = p.add_run(desc)
    run2.font.size  = Pt(11)
    run2.font.color.rgb = C_BODY
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 4. REQUIREMENTS
# ═════════════════════════════════════════════════════════════
add_h1(doc, "4. REQUIREMENTS")
add_body(doc,
    "The following section details both the software and hardware requirements needed to develop "
    "and deploy GoalFit AI Pro."
)
doc.add_paragraph()


# 4.1
add_h2(doc, "4.1 Software Requirements")
add_h3(doc, "Server Side")
header_table(doc,
    ["Component", "Tool / Library", "Version / Notes"],
    [
        ("Programming Language",   "Python",              "3.12+"),
        ("Web Framework",          "Flask",               "3.x"),
        ("Database",               "MySQL",               "8.0+"),
        ("DB Connector",           "mysql-connector-python","8.x"),
        ("ORM / Query",            "Raw SQL (cursor)",    "–"),
        ("Authentication",         "Werkzeug (password hash)", "3.x"),
        ("Rate Limiting",          "Flask-Limiter",       "3.x"),
        ("PDF Generation",         "ReportLab / fpdf2",   "Latest"),
        ("Env Config",             "python-dotenv",       "1.x"),
        ("Template Engine",        "Jinja2",              "Bundled with Flask"),
        ("Payment Gateway",        "Razorpay Python SDK", "Latest"),
        ("Doc Generation",         "python-docx",         "1.2+"),
    ],
    col_widths=[1.8, 2.5, 2.0]
)
doc.add_paragraph()

add_h3(doc, "Client Side")
header_table(doc,
    ["Component", "Technology", "Notes"],
    [
        ("Markup",        "HTML5",                  "Semantic layout"),
        ("Styling",       "CSS3 + Bootstrap 5",     "Custom professional.css dark theme"),
        ("Scripting",     "Vanilla JavaScript",     "No frontend framework"),
        ("Charts",        "Chart.js",               "Revenue, weight, donut charts"),
        ("Icons",         "Lucide Icons",           "SVG icon set"),
        ("Animations",    "AOS (Animate on Scroll)","Scroll-triggered transitions"),
        ("Fonts",         "Google Fonts – Outfit",  "Typography"),
        ("HTTP Client",   "Fetch API",              "Chat messages, async updates"),
    ],
    col_widths=[1.8, 2.5, 2.0]
)
doc.add_paragraph()

add_h3(doc, "Development Tools")
header_table(doc,
    ["Tool", "Purpose"],
    [
        ("Visual Studio Code",  "Primary code editor"),
        ("Git + GitHub",        "Version control & collaboration"),
        ("MySQL Workbench",     "Database design & management"),
        ("Postman",             "API testing"),
        ("Windows 11",         "Development OS"),
        ("Chrome DevTools",     "Frontend debugging"),
    ],
    col_widths=[2.5, 4.0]
)
doc.add_paragraph()


# 4.2
add_h2(doc, "4.2 Hardware Requirements")
add_h3(doc, "Server Side (Deployment)")
header_table(doc,
    ["Component", "Minimum", "Recommended"],
    [
        ("Processor",   "Intel Core i3 / AMD Ryzen 3",   "Intel Core i5+ / Ryzen 5+"),
        ("RAM",         "4 GB",                          "8 GB+"),
        ("Storage",     "20 GB SSD",                     "50 GB SSD"),
        ("Network",     "10 Mbps",                       "100 Mbps+"),
        ("OS",          "Ubuntu 20.04 LTS",              "Ubuntu 22.04 LTS"),
    ],
    col_widths=[1.8, 2.5, 2.3]
)
doc.add_paragraph()

add_h3(doc, "Client Side (User Device)")
header_table(doc,
    ["Component", "Minimum Requirement"],
    [
        ("Device",       "Desktop, Laptop, Tablet, Smartphone"),
        ("Browser",      "Chrome 90+, Firefox 85+, Edge 90+, Safari 14+"),
        ("RAM",          "2 GB"),
        ("Screen",       "1280 × 720 or higher"),
        ("Internet",     "Stable broadband / 4G connection"),
    ],
    col_widths=[2.0, 4.5]
)
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 5. TIME DURATION
# ═════════════════════════════════════════════════════════════
add_h1(doc, "5. TIME DURATION")
add_body(doc, "Estimated total development time: approximately 5–6 months.\n")
header_table(doc,
    ["Phase", "Activities", "Duration"],
    [
        ("Phase 1 – Planning & Analysis",
         "Requirement gathering, ER design, system flow",
         "2 weeks"),
        ("Phase 2 – Database & Backend Setup",
         "MySQL schema creation, Flask app skeleton, blueprints, auth",
         "3 weeks"),
        ("Phase 3 – Core User Features",
         "Health profile, BMI, AI diet/workout plans, water tracker",
         "4 weeks"),
        ("Phase 4 – Professional Marketplace",
         "Professional registration, marketplace listings, hire flow, Razorpay",
         "3 weeks"),
        ("Phase 5 – Professional SaaS Portal",
         "Pro dashboard, client management, meal/workout builders, chat",
         "4 weeks"),
        ("Phase 6 – Admin Panel",
         "Admin dashboard, user/professional management, payments overview",
         "2 weeks"),
        ("Phase 7 – Progress & Tracking Features",
         "Progress photos, weight logs, notifications, PDF generator",
         "2 weeks"),
        ("Phase 8 – Testing & Deployment",
         "Unit testing, browser testing, bug fixes, seed data, deployment",
         "2 weeks"),
    ],
    col_widths=[2.2, 3.5, 1.0]
)
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 6. TECHNOLOGY USED
# ═════════════════════════════════════════════════════════════
add_h1(doc, "6. TECHNOLOGY USED")
tech_blocks = {
    "Python 3.12": (
        "Primary backend language. Used for all server-side logic including routing, "
        "database interaction, PDF generation, and authentication."
    ),
    "Flask (Micro Web Framework)": (
        "Lightweight WSGI web application framework. Uses the Blueprint pattern to "
        "separate concerns: auth, dashboard, marketplace, professional portal, admin, etc. "
        "Provides Jinja2 templating, session management, and request handling."
    ),
    "MySQL 8.0": (
        "Relational database. Stores 33 tables covering users, professionals, assignments, "
        "meals, workouts, payments, chat messages, notifications, and more."
    ),
    "Jinja2 Templating": (
        "Server-side rendering of HTML pages with dynamic data injection. Used across "
        "all 50+ templates in the project."
    ),
    "Bootstrap 5 + Custom CSS": (
        "Responsive grid system for layout. Custom dark-themed CSS (professional.css, style.css, "
        "admin.css) provides premium glassmorphism and neon-accent design."
    ),
    "Chart.js": (
        "Interactive JavaScript chart library. Used for revenue line charts, client "
        "weight progress charts, and client distribution donut charts."
    ),
    "Razorpay Payment Gateway": (
        "Handles secure online payments when users hire professionals. Razorpay order IDs "
        "and payment verifications are processed server-side."
    ),
    "ReportLab / FPDF2": (
        "Generates downloadable, formatted PDF files for personalised meal plans "
        "and workout routines."
    ),
    "Flask-Limiter": (
        "Rate limiting middleware to prevent abuse. Default limits: 200 requests/day, "
        "50 requests/hour."
    ),
    "Werkzeug Security": (
        "Password hashing (PBKDF2-SHA256) and verification for both user and professional "
        "accounts."
    ),
    "python-dotenv": (
        "Loads environment variables from .env file (DB credentials, secret keys), "
        "separating configuration from code."
    ),
    "Lucide Icons + AOS": (
        "Lucide provides a clean SVG icon set. AOS (Animate on Scroll) library delivers "
        "smooth scroll-triggered CSS animations."
    ),
}
for tech, desc in tech_blocks.items():
    p = doc.add_paragraph()
    r1 = p.add_run(f"{tech}: ")
    r1.font.bold  = True
    r1.font.size  = Pt(11)
    r1.font.color.rgb = C_SUBHEAD
    r2 = p.add_run(desc)
    r2.font.size  = Pt(11)
    r2.font.color.rgb = C_BODY
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 7. FUNCTIONAL AND NON-FUNCTIONAL REQUIREMENTS
# ═════════════════════════════════════════════════════════════
add_h1(doc, "7. FUNCTIONAL AND NON-FUNCTIONAL REQUIREMENTS")
add_h2(doc, "7.1 Functional Requirements")
functional = [
    ("FR-01", "User Registration & Login",    "Users can register, log in, and log out securely. Passwords are hashed."),
    ("FR-02", "Health Profile Setup",         "First-time users fill age, gender, height, weight, goal type, and diet preference."),
    ("FR-03", "BMI Calculation",              "Computes BMI and categorises it (underweight, normal, overweight, obese)."),
    ("FR-04", "AI Diet Plan Generation",      "Personalised meal plan generated based on goal type and diet preference (vegetarian, non-vegetarian, vegan)."),
    ("FR-05", "AI Workout Plan Generation",   "Workout plan assigned by muscle group, difficulty level, and goal type."),
    ("FR-06", "Goal Prediction",              "System calculates estimated weeks and target date to achieve goal weight."),
    ("FR-07", "Step Recommendation",          "Daily steps, calories to burn, and distance target recommended per user profile."),
    ("FR-08", "Water Intake Tracker",         "Users log daily water consumption; visual progress bar shown."),
    ("FR-09", "Progress Logging",             "Users log weight periodically; weight chart displayed."),
    ("FR-10", "Progress Photo Gallery",       "Users upload progress photos with sharing controls."),
    ("FR-11", "Professional Marketplace",     "Browse trainers/dieticians with cards showing name, rating, and pricing."),
    ("FR-12", "Professional Profile",         "Detailed profile page with bio, transformations, pricing plans, and hire button."),
    ("FR-13", "Hire Professional",            "User selects plan, pays via Razorpay; hire request created."),
    ("FR-14", "Chat System",                  "Real-time overlay chat between professional and client."),
    ("FR-15", "Custom Diet Plan (Pro)",       "Professionals create custom meal library and assign meal plans to clients."),
    ("FR-16", "Custom Workout Plan (Pro)",    "Professionals create exercises and assign day-wise workout plans to clients."),
    ("FR-17", "Transformation Gallery",       "Professionals upload client transformation stories (before/after weight, description)."),
    ("FR-18", "Notifications",               "Professionals receive notifications for hire requests, payments, and reviews."),
    ("FR-19", "PDF Report Download",          "Users download personalised diet and workout plans as formatted PDFs."),
    ("FR-20", "Admin Dashboard",              "Admin manages users, professionals (verify/suspend), views payments and feedback."),
]
header_table(doc,
    ["ID", "Feature", "Description"],
    functional,
    col_widths=[0.7, 1.9, 4.0]
)
doc.add_paragraph()

add_h2(doc, "7.2 Non-Functional Requirements")
nfr = [
    ("NFR-01", "Performance",    "Page load time < 2 seconds on a standard broadband connection."),
    ("NFR-02", "Security",       "Passwords hashed with PBKDF2-SHA256; CSRF protection via Flask session tokens; rate limiting."),
    ("NFR-03", "Scalability",    "Blueprint architecture allows independent module scaling."),
    ("NFR-04", "Usability",      "Dark-themed, mobile-responsive design with accessibility considerations."),
    ("NFR-05", "Reliability",    "MySQL foreign key constraints and transaction commits ensure data integrity."),
    ("NFR-06", "Maintainability","Modular codebase with 27+ blueprints, each in a separate file."),
    ("NFR-07", "Portability",    "Deployable on any Linux/Windows server with Python 3.12 and MySQL 8."),
    ("NFR-08", "Availability",   "Target uptime of 99.5% in production environment."),
]
header_table(doc,
    ["ID", "Category", "Requirement"],
    nfr,
    col_widths=[0.7, 1.5, 4.4]
)
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 8. DESIGN
# ═════════════════════════════════════════════════════════════
add_h1(doc, "8. DESIGN")

# 8.1
add_h2(doc, "8.1 System Flow Diagram")
add_body(doc,
    "The following describes the high-level system flow for GoalFit AI Pro:"
)
flow_steps = [
    ("1", "User visits the platform",        "Landing page → Login / Register"),
    ("2", "First-time login",               "Redirect to Health Profile setup form"),
    ("3", "Dashboard",                       "AI generates personalised diet & workout plan; displays BMI, goals, water tracker"),
    ("4", "Marketplace",                     "User browses trainers/dieticians, views profiles, selects plan"),
    ("5", "Payment",                         "Razorpay checkout; on success, hire_request created with status=accepted"),
    ("6", "Professional Portal",             "Professional receives notification, accepts client, assigns custom plans"),
    ("7", "Client views custom plans",       "User dashboard updates with professional's custom diet/workout plans"),
    ("8", "Progress Tracking",              "User logs weight, meals, workouts; charts and gallery updated"),
    ("9", "PDF Download",                   "User exports personalised report as PDF"),
    ("10","Admin oversight",               "Admin verifies professionals, manages users, views platform analytics"),
]
header_table(doc,
    ["Step", "Actor", "Action"],
    flow_steps,
    col_widths=[0.5, 1.8, 4.3]
)
doc.add_paragraph()


# 8.2
add_h2(doc, "8.2 UML Diagrams / Data Flow Diagrams")
add_body(doc,
    "This section presents the key UML diagrams describing the system behaviour and data flow "
    "of GoalFit AI Pro. Each diagram is described textually below; visual diagrams should be "
    "inserted at these locations."
)
doc.add_paragraph()

# 8.2.1
add_h3(doc, "8.2.1 Use Case Diagram (Context / 0 Level)")
add_body(doc, "Actors:")
for a in ["Guest User", "Registered User", "Trainer", "Dietician", "Admin"]:
    add_bullet(doc, a)
add_body(doc, "\nKey Use Cases:")
use_cases = [
    ("Guest User",        ["Register", "Login", "Browse Marketplace"]),
    ("Registered User",   ["Setup Health Profile", "View Diet Plan", "View Workout Plan",
                            "Log Progress", "Track Water", "Hire Professional",
                            "Download PDF", "Chat with Professional", "View Predictions"]),
    ("Trainer",           ["Register as Professional", "Login to Pro Portal",
                            "Manage Clients", "Create Exercise Library",
                            "Assign Workout Plans", "Upload Transformations",
                            "View Earnings"]),
    ("Dietician",         ["Create Meal Library", "Assign Diet Plans",
                            "View Client Progress", "Receive Notifications"]),
    ("Admin",             ["Verify Professionals", "Manage Users", "View Payments",
                            "Reply to Feedback", "View Platform Analytics"]),
]
for actor, cases in use_cases:
    p = doc.add_paragraph()
    r = p.add_run(f"{actor}: ")
    r.font.bold  = True
    r.font.color.rgb = C_SUBHEAD
    r.font.size  = Pt(11)
    doc.add_paragraph(", ".join(cases)).runs[0].font.size = Pt(11)
doc.add_paragraph()


# 8.2.2
add_h3(doc, "8.2.2 Activity Diagram (First Level)")
add_body(doc,
    "Activity diagram for the User Registration and Health Profile Setup workflow:"
)
steps = [
    "Start",
    "User opens registration form → enters name, email, password",
    "System validates email uniqueness → password hashed → record inserted in users table",
    "User redirected to Login page",
    "User logs in → session created",
    "System checks: is this a first_time_login? → YES",
    "Redirect to Health Profile form",
    "User enters age, gender, height, weight, goal, activity level, diet preference",
    "System saves to user_health table",
    "System calculates BMI → saves to bmi_records",
    "System runs Goal Prediction algorithm → saves to goal_predictions",
    "System generates Step Recommendation → saves to step_recommendations",
    "Redirect to Dashboard → Display personalised plan",
    "End",
]
for s in steps:
    add_bullet(doc, s)
doc.add_paragraph()


# 8.2.3
add_h3(doc, "8.2.3 Sequence Diagram (Second Level)")
add_body(doc,
    "Sequence for Hiring a Professional via Marketplace:"
)
sequence = [
    ("User",         "Browser",       "Clicks 'View Trainers' on Marketplace"),
    ("Browser",      "Flask /marketplace/trainers", "GET request"),
    ("Flask",        "MySQL",         "SELECT * FROM professionals WHERE role='trainer'"),
    ("MySQL",        "Flask",         "Returns professionals list"),
    ("Flask",        "Browser",       "Renders marketplace.html with trainer cards"),
    ("User",         "Browser",       "Clicks trainer card → View Profile"),
    ("Browser",      "Flask /professional/<id>","GET request"),
    ("Flask",        "MySQL",         "Fetch professional + pricing + reviews + transformations"),
    ("Flask",        "Browser",       "Renders profile page"),
    ("User",         "Browser",       "Clicks 'Hire Now' → selects plan"),
    ("Browser",      "Razorpay",      "Creates payment order via /payment/create-order"),
    ("Razorpay",     "Browser",       "Returns order_id"),
    ("User",         "Razorpay",      "Completes payment"),
    ("Browser",      "Flask /payment/verify","POST with payment_id"),
    ("Flask",        "MySQL",         "INSERT hire_request (status=accepted)"),
    ("Flask",        "MySQL",         "INSERT payment record"),
    ("Flask",        "MySQL",         "INSERT notification for professional"),
    ("Flask",        "Browser",       "Redirect to success page"),
]
header_table(doc,
    ["From", "To", "Message / Action"],
    sequence,
    col_widths=[1.5, 2.0, 3.1]
)
doc.add_paragraph()


# 8.2.4
add_h3(doc, "8.2.4 Class Diagram")
add_body(doc,
    "Key classes / entities and their relationships in GoalFit AI Pro:"
)
classes = [
    ("User",             "id, name, email, password, role, created_at",
                         "Health, BMI, WorkoutLog, MealLog, WaterLog, Progress, HireRequest"),
    ("UserHealth",       "id, user_id, age, gender, height_cm, weight_kg, target_weight, goal_type",
                         "User (1:1)"),
    ("Professional",     "id, full_name, email, role, bio, experience_years, rating",
                         "Pricing, ClientAssignment, Transformation, Review"),
    ("HireRequest",      "id, user_id, professional_id, plan_type, status, payment_status",
                         "User (N:1), Professional (N:1), Payment (1:1)"),
    ("ClientAssignment", "id, user_id, professional_id, plan_type, start_date, end_date, status",
                         "CustomDietPlan, CustomWorkoutPlan"),
    ("CustomDietPlan",   "id, user_id, professional_id, plan_name, goal, notes",
                         "CustomDietPlanMeal (1:N)"),
    ("CustomWorkoutPlan","id, user_id, professional_id, plan_name, goal, notes",
                         "CustomWorkoutPlanExercise (1:N)"),
    ("Payment",          "id, user_id, professional_id, razorpay_payment_id, amount",
                         "HireRequest (1:1)"),
    ("Notification",     "id, professional_id, user_id, type, message, is_read",
                         "Professional (N:1)"),
    ("ChatMessage",      "id, sender_id, sender_role, receiver_id, receiver_role, message",
                         "User / Professional (polymorphic)"),
]
header_table(doc,
    ["Class / Entity", "Key Attributes", "Relationships"],
    classes,
    col_widths=[1.7, 2.8, 2.1]
)
doc.add_paragraph()


# 8.3
add_h2(doc, "8.3 ER Diagram")
add_body(doc,
    "The Entity-Relationship (ER) diagram for GoalFit AI Pro includes the following "
    "primary entities and their relationships:"
)
er_entities = [
    ("users",                 "id (PK), name, email, password, role, created_at",       "Central entity linking to all user data"),
    ("user_health",           "id (PK), user_id (FK), age, gender, height_cm, weight_kg, goal_type", "1:1 with users"),
    ("bmi_records",           "id (PK), user_id (FK), bmi_value, bmi_category, recorded_date",       "N:1 with users"),
    ("progress_logs",         "id (PK), user_id (FK), weight_kg, log_date",                          "N:1 with users"),
    ("professionals",         "id (PK), full_name, email, role, bio, experience_years, rating",       "Central professional entity"),
    ("professional_pricing",  "id (PK), professional_id (FK), plan_type, duration_days, price",       "N:1 with professionals"),
    ("hire_requests",         "id (PK), user_id (FK), professional_id (FK), plan_type, status",       "N:N bridge users ↔ professionals"),
    ("client_assignments",    "id (PK), user_id (FK), professional_id (FK), status",                  "Active coaching relationships"),
    ("custom_diet_plans",     "id (PK), user_id (FK), professional_id (FK), plan_name, goal",        "N:1 users, N:1 professionals"),
    ("custom_workout_plans",  "id (PK), user_id (FK), professional_id (FK), plan_name, goal",        "N:1 users, N:1 professionals"),
    ("professional_meals",    "id (PK), professional_id (FK), meal_name, calories, protein",         "N:1 professionals"),
    ("professional_workouts", "id (PK), professional_id (FK), workout_name, target_muscle, sets",    "N:1 professionals"),
    ("transformations",       "id (PK), professional_id (FK), client_name, before_weight, after_weight","N:1 professionals"),
    ("payments",              "id (PK), user_id (FK), professional_id (FK), razorpay_payment_id",    "N:N bridge"),
    ("chat_messages",         "id (PK), sender_id, sender_role, receiver_id, receiver_role, message","Polymorphic messaging"),
    ("notifications",         "id (PK), professional_id (FK), user_id (FK), type, is_read",          "N:1 professionals / users"),
    ("professional_reviews",  "id (PK), professional_id (FK), user_id (FK), rating, review_text",    "N:N bridge"),
    ("water_logs",            "id (PK), user_id (FK), glasses, goal_glasses, log_date",              "N:1 users"),
    ("diet_logs",             "id (PK), user_id (FK), meal_id (FK), log_date, is_completed",         "N:1 users, N:1 diet_meals"),
    ("progress_photos",       "id (PK), user_id (FK), photo_path, log_date, is_shared",              "N:1 users"),
]
header_table(doc,
    ["Entity / Table", "Key Columns", "Relationship Notes"],
    er_entities,
    col_widths=[1.8, 2.7, 2.1]
)
doc.add_paragraph()


# 8.4 DATA DICTIONARY
add_h2(doc, "8.4 Data Dictionary")
add_body(doc, "Selected tables and their column descriptions are given below.\n")

# users
add_h3(doc, "Table: users")
header_table(doc,
    ["Column", "Data Type", "Constraint", "Description"],
    [
        ("id",         "INT",          "PK, AUTO_INCREMENT", "Unique user identifier"),
        ("name",       "VARCHAR(100)", "NOT NULL",           "Full display name"),
        ("email",      "VARCHAR(100)", "UNIQUE, NOT NULL",   "Login email address"),
        ("password",   "VARCHAR(255)", "NOT NULL",           "Hashed password (Werkzeug)"),
        ("role",       "VARCHAR(20)",  "DEFAULT 'user'",     "user | admin"),
        ("created_at", "TIMESTAMP",    "DEFAULT NOW()",      "Account creation timestamp"),
    ],
    col_widths=[1.3, 1.3, 1.8, 2.2]
)
doc.add_paragraph()

# professionals
add_h3(doc, "Table: professionals")
header_table(doc,
    ["Column", "Data Type", "Constraint", "Description"],
    [
        ("id",              "INT",          "PK, AUTO_INCREMENT",  "Unique professional ID"),
        ("full_name",       "VARCHAR(100)", "NOT NULL",            "Full name"),
        ("email",           "VARCHAR(100)", "UNIQUE",              "Login email"),
        ("password",        "VARCHAR(255)", "NOT NULL",            "Hashed password"),
        ("phone",           "VARCHAR(20)",  "–",                   "Contact number"),
        ("role",            "ENUM",         "trainer|dietician|both","Professional role"),
        ("bio",             "TEXT",         "–",                   "Biography / intro"),
        ("experience_years","INT",          "–",                   "Years of experience"),
        ("specialization",  "VARCHAR(255)", "–",                   "Areas of expertise"),
        ("is_verified",     "BOOLEAN",      "DEFAULT FALSE",       "Admin-verified flag"),
        ("rating",          "FLOAT",        "DEFAULT 0.0",         "Average client rating"),
    ],
    col_widths=[1.4, 1.2, 1.6, 2.4]
)
doc.add_paragraph()

# hire_requests
add_h3(doc, "Table: hire_requests")
header_table(doc,
    ["Column", "Data Type", "Constraint", "Description"],
    [
        ("id",              "INT",         "PK",              "Unique request ID"),
        ("user_id",         "INT",         "FK → users",      "Client placing the request"),
        ("professional_id", "INT",         "FK → professionals","Target professional"),
        ("plan_type",       "VARCHAR(50)", "–",               "Selected plan type"),
        ("goal_type",       "VARCHAR(50)", "–",               "User fitness goal"),
        ("payment_status",  "VARCHAR(50)", "DEFAULT 'pending'","pending | paid | failed"),
        ("status",          "ENUM",        "pending|accepted|rejected|completed","Request lifecycle"),
        ("created_at",      "TIMESTAMP",   "DEFAULT NOW()",   "Timestamp"),
    ],
    col_widths=[1.4, 1.2, 1.6, 2.4]
)
doc.add_paragraph()

# payments
add_h3(doc, "Table: payments")
header_table(doc,
    ["Column", "Data Type", "Constraint", "Description"],
    [
        ("id",                  "INT",         "PK",              "Payment record ID"),
        ("user_id",             "INT",         "FK → users",      "Paying user"),
        ("professional_id",     "INT",         "FK → professionals","Recipient professional"),
        ("hire_request_id",     "INT",         "FK → hire_requests","Associated hire request"),
        ("razorpay_payment_id", "VARCHAR(100)","UNIQUE",          "Razorpay transaction ID"),
        ("amount",              "FLOAT",       "–",               "Total amount paid (INR)"),
        ("commission_amount",   "FLOAT",       "–",               "Platform commission (15%)"),
        ("professional_amount", "FLOAT",       "–",               "Amount after commission"),
        ("payment_status",      "VARCHAR(50)", "DEFAULT 'pending'","pending | paid | failed"),
        ("created_at",          "TIMESTAMP",   "DEFAULT NOW()",   "Payment timestamp"),
    ],
    col_widths=[1.6, 1.1, 1.5, 2.4]
)
doc.add_paragraph()

doc.add_paragraph(
    "[NOTE: Screenshots of all application modules should be inserted in Section 8.5. "
    "Please capture and embed screenshots of the following pages: Login, Registration, "
    "Health Profile, User Dashboard, Diet Plan, Workout Plan, Water Tracker, "
    "Progress Gallery, Marketplace, Professional Profile, Hire Flow, Payment, "
    "Professional Dashboard, Client Management, Diet Plan Builder, Workout Plan Builder, "
    "Transformations, Reviews, Notifications, Settings, Admin Dashboard, "
    "Admin Users, Admin Professionals, Admin Payments, PDF Report.]"
).runs[0].font.italic = True


# ═════════════════════════════════════════════════════════════
# 8.5  (placeholder)
# ═════════════════════════════════════════════════════════════
add_h2(doc, "8.5 Module Screenshots")
add_body(doc,
    "Screenshots of each module are to be inserted here. Refer to the NOTE above for "
    "the list of required screenshots. Each screenshot should include a caption "
    "describing the module it represents."
)
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 9. APPLICATION
# ═════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, "9. APPLICATION")
app_areas = [
    ("Personal Health Management",
     "Individuals who want a structured and AI-driven approach to managing their weight, "
     "diet, and fitness routine without needing a gym membership."),
    ("Professional Fitness Coaching",
     "Certified trainers and dieticians who need a digital workspace to onboard clients, "
     "create personalised plans, and track client outcomes."),
    ("Health-Conscious Workplaces",
     "Corporate wellness programs can subscribe employees to the platform for group "
     "health monitoring and goal tracking."),
    ("Medical Nutrition Therapy",
     "Registered dieticians can use the platform to manage patients requiring therapeutic "
     "diets (e.g., PCOS, diabetes, weight-related conditions)."),
    ("Fitness Education",
     "Students and trainees in nutrition science and sports medicine can use the platform "
     "to simulate client management scenarios."),
    ("E-Commerce for Fitness Services",
     "The marketplace enables trainers and dieticians to monetise their expertise through "
     "a structured, payment-enabled digital storefront."),
]
for title, desc in app_areas:
    p = doc.add_paragraph()
    r1 = p.add_run(f"• {title}: ")
    r1.font.bold  = True
    r1.font.size  = Pt(11)
    r1.font.color.rgb = C_SUBHEAD
    r2 = p.add_run(desc)
    r2.font.size  = Pt(11)
    r2.font.color.rgb = C_BODY
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 10. EXPECTED OUTCOMES
# ═════════════════════════════════════════════════════════════
add_h1(doc, "10. EXPECTED OUTCOMES")
outcomes = [
    "Users will be able to generate AI-personalised diet and workout plans within minutes of registration.",
    "Users will achieve improved awareness of their BMI, calorie targets, daily steps, and progress over time.",
    "Trainers and dieticians will acquire an end-to-end digital workspace to manage clients, create plans, upload results, and earn revenue.",
    "The marketplace will connect fitness professionals with potential clients, increasing earning opportunities by 30–40% compared to offline-only models.",
    "Admin will gain real-time visibility into platform health—user counts, professional counts, total revenue, and feedback.",
    "The PDF export feature will increase plan adherence by giving users a tangible, offline-accessible document.",
    "The notifications and chat system will improve professional–client communication and engagement.",
    "The platform's modular architecture will allow future AI enhancements (e.g., computer vision for posture correction, LLM-based coaching) without restructuring the codebase.",
    "GoalFit AI Pro is expected to serve as a reference implementation of a production-grade Flask SaaS product for academic and entrepreneurial purposes.",
]
for o in outcomes:
    add_bullet(doc, o)
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 11. FUTURE SCOPE
# ═════════════════════════════════════════════════════════════
add_h1(doc, "11. FUTURE SCOPE")
future = [
    ("AI Chatbot Coaching",
     "Integrate an LLM (Gemini / GPT-4) powered coaching chatbot that can answer "
     "fitness and nutrition questions in real-time, 24/7."),
    ("Computer Vision – Posture Analysis",
     "Use TensorFlow.js or MediaPipe to analyse exercise posture via webcam and "
     "provide real-time corrective feedback."),
    ("Food Image Recognition",
     "Allow users to photograph their meals; AI identifies food items and automatically "
     "logs calories and macros."),
    ("Wearable Device Integration",
     "Connect with fitness wearables (Apple Watch, Fitbit, Google Fit) to automatically "
     "sync steps, heart rate, and sleep data."),
    ("Mobile App (React Native / Flutter)",
     "Package GoalFit AI Pro as a cross-platform mobile application for iOS and Android "
     "with push notifications."),
    ("Social & Community Features",
     "Add fitness challenges, leaderboards, community forums, and peer-to-peer progress sharing."),
    ("Subscription Model",
     "Introduce tiered monthly subscriptions (Free, Premium, Pro) with feature gating "
     "and automated billing via Razorpay Subscriptions."),
    ("Multi-Language Support",
     "Localise the platform in Hindi, Gujarati, and other Indian regional languages "
     "to increase accessibility."),
    ("Telehealth Integration",
     "Enable video consultation appointments between users and professionals inside the platform."),
    ("Advanced Analytics for Professionals",
     "Provide professionals with client retention metrics, churn predictions, and revenue forecasts."),
]
for title, desc in future:
    p = doc.add_paragraph()
    r1 = p.add_run(f"• {title}: ")
    r1.font.bold  = True
    r1.font.size  = Pt(11)
    r1.font.color.rgb = C_SUBHEAD
    r2 = p.add_run(desc)
    r2.font.size  = Pt(11)
    r2.font.color.rgb = C_BODY
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════
# 12. REFERENCES
# ═════════════════════════════════════════════════════════════
add_h1(doc, "12. REFERENCES")
references = [
    "[1]  Flask Documentation – https://flask.palletsprojects.com/",
    "[2]  MySQL 8.0 Reference Manual – https://dev.mysql.com/doc/",
    "[3]  Bootstrap 5 Documentation – https://getbootstrap.com/docs/5.3/",
    "[4]  Chart.js Documentation – https://www.chartjs.org/docs/latest/",
    "[5]  Razorpay Payment Gateway – https://razorpay.com/docs/",
    "[6]  ReportLab PDF Toolkit – https://www.reportlab.com/documentation/",
    "[7]  Werkzeug Security Utilities – https://werkzeug.palletsprojects.com/en/latest/utils/",
    "[8]  Flask-Limiter – https://flask-limiter.readthedocs.io/en/stable/",
    "[9]  python-dotenv – https://pypi.org/project/python-dotenv/",
    "[10] Lucide Icons – https://lucide.dev/",
    "[11] AOS – Animate on Scroll Library – https://michalsnik.github.io/aos/",
    "[12] Jinja2 Templating Engine – https://jinja.palletsprojects.com/",
    "[13] WHO – Physical Activity Guidelines – https://www.who.int/news-room/fact-sheets/detail/physical-activity",
    "[14] Harvard Health – Diet and Nutrition – https://www.health.harvard.edu/topics/diet-and-weight-loss",
    "[15] NSCA – National Strength & Conditioning Association – https://www.nsca.com/",
]
for ref in references:
    p = doc.add_paragraph(ref)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = C_BODY
doc.add_paragraph()

# ─── Save ───────────────────────────────────────────────────
os.makedirs("documentation", exist_ok=True)
doc.save(DOC_PATH)
print(f"Documentation saved: {DOC_PATH}")
